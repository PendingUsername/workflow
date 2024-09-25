import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import RGBColor
import difflib
import os

try:
    import win32com.client as win32
except ImportError:
    messagebox.showerror("Import Error", "Please install pywin32 by running: pip install pywin32")


def get_text_from_docx(docx_path):
    """Extract all text from a .docx Word document."""
    doc = Document(docx_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return full_text


def get_text_from_doc(doc_path):
    """Extract all text from a .doc Word document using pywin32."""
    # Normalize the file path
    doc_path = os.path.normpath(doc_path)
    
    if not os.path.exists(doc_path):
        messagebox.showerror("File Error", f"File not found: {doc_path}")
        return None

    # Open the document using Word COM interface
    word = win32.Dispatch("Word.Application")
    try:
        doc = word.Documents.Open(doc_path)
        full_text = doc.Content.Text.split("\r")
        doc.Close(False)
        return full_text
    except Exception as e:
        messagebox.showerror("Error", f"Error opening the .doc file: {e}")
        return None
    finally:
        word.Quit()


def get_text_from_file(file_path):
    """Determine the file type and extract text accordingly."""
    _, file_extension = os.path.splitext(file_path)
    if file_extension.lower() == ".docx":
        return get_text_from_docx(file_path)
    elif file_extension.lower() == ".doc":
        return get_text_from_doc(file_path)
    else:
        messagebox.showerror("Unsupported Format", "Please select a .docx or .doc file.")
        return None


def highlight_differences(doc1, doc2, differences):
    """Highlight the differences in a new .docx document."""
    output_doc = Document()

    # Create lists of text from each document
    diff = difflib.ndiff(doc1, doc2)

    for line in diff:
        if line.startswith('+ '):
            # Add added text in green
            paragraph = output_doc.add_paragraph(line[2:])
            if paragraph.runs:  # Ensure the run exists
                run = paragraph.runs[0]
                run.font.color.rgb = RGBColor(0, 255, 0)  # Green for additions
        elif line.startswith('- '):
            # Add deleted text in red
            paragraph = output_doc.add_paragraph(line[2:])
            if paragraph.runs:  # Ensure the run exists
                run = paragraph.runs[0]
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red for deletions
        else:
            # Keep unchanged text
            output_doc.add_paragraph(line[2:])

    # Save the new document
    save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                             filetypes=[("Word Documents", "*.docx")],
                                             title="Save Highlighted Differences")
    if save_path:
        output_doc.save(save_path)
        messagebox.showinfo("Success", f"Differences saved in {save_path}")


def select_file1():
    file_path = filedialog.askopenfilename(
        title="Select the first Word document",
        filetypes=(("Word files", "*.docx *.doc"), ("All files", "*.*"))
    )
    if file_path:
        entry1.delete(0, tk.END)
        entry1.insert(0, file_path)


def select_file2():
    file_path = filedialog.askopenfilename(
        title="Select the second Word document",
        filetypes=(("Word files", "*.docx *.doc"), ("All files", "*.*"))
    )
    if file_path:
        entry2.delete(0, tk.END)
        entry2.insert(0, file_path)


def compare_action():
    docx_path1 = entry1.get()
    docx_path2 = entry2.get()

    if not docx_path1 or not docx_path2:
        messagebox.showerror("Error", "Please select both documents.")
        return

    text1 = get_text_from_file(docx_path1)
    text2 = get_text_from_file(docx_path2)

    if text1 is None or text2 is None:
        return

    highlight_differences(text1, text2, None)


# Set up the main application window
root = tk.Tk()
root.title("Word Document Comparator")

# File selection inputs
frame = tk.Frame(root)
frame.pack(padx=10, pady=10)

label1 = tk.Label(frame, text="Select the first document:")
label1.grid(row=0, column=0, padx=5, pady=5, sticky="e")

entry1 = tk.Entry(frame, width=50)
entry1.grid(row=0, column=1, padx=5, pady=5)

button1 = tk.Button(frame, text="Browse", command=select_file1)
button1.grid(row=0, column=2, padx=5, pady=5)

label2 = tk.Label(frame, text="Select the second document:")
label2.grid(row=1, column=0, padx=5, pady=5, sticky="e")

entry2 = tk.Entry(frame, width=50)
entry2.grid(row=1, column=1, padx=5, pady=5)

button2 = tk.Button(frame, text="Browse", command=select_file2)
button2.grid(row=1, column=2, padx=5, pady=5)

# Compare button
compare_button = tk.Button(root, text="Compare", command=compare_action)
compare_button.pack(pady=10)

# Start the GUI event loop
root.mainloop()

