# Run the program using python3 .\exam_ready.py
import tkinter as tk
from tkinter import simpledialog, messagebox
from docx import Document

def generate_exam_email():
    # Create a Tkinter root window and hide it
    root = tk.Tk()
    root.withdraw()

    # Prompt the user for input
    exam_name = simpledialog.askstring("Input", "Enter the exam name:", parent=root)
    mcq_exam_name = simpledialog.askstring("Input", "Enter the MCQ exam file name:", parent=root)
    essay_exam_name = simpledialog.askstring("Input", "Enter the essay exam file name:", parent=root)
    exam_date = simpledialog.askstring("Input", "Enter the exam date (e.g., Friday, August 30, 2024):", parent=root)
    exam_time = simpledialog.askstring("Input", "Enter the exam time (e.g., 8:00 AM):", parent=root)
    mcq_start_time = simpledialog.askstring("Input", "Enter the MCQ start time (e.g., 8:00 AM):", parent=root)
    mcq_end_time = simpledialog.askstring("Input", "Enter the MCQ end time (e.g., 10:30 AM):", parent=root)
    mcq_duration = simpledialog.askstring("Input", "Enter the MCQ duration (e.g., 150 mins):", parent=root)
    essay_start_time = simpledialog.askstring("Input", "Enter the essay start time (e.g., 10:30 AM):", parent=root)
    essay_end_time = simpledialog.askstring("Input", "Enter the essay end time (e.g., 11:45 AM):", parent=root)
    essay_duration = simpledialog.askstring("Input", "Enter the essay duration (e.g., 75 mins):", parent=root)
    exam_location = simpledialog.askstring("Input", "Enter the exam location (e.g., Room 120):", parent=root)


    # Create a Word document
    doc = Document()

    # Add content to the document based on user input
    doc.add_paragraph(f"Good afternoon, all.")
    doc.add_paragraph("")
    doc.add_paragraph(f"The {exam_name} exam(s) are ready for you to download. Please make sure your exams are downloaded by Thursday so that you are prepared for testing on Friday morning. The file download names are \"{mcq_exam_name}\" and \"{essay_exam_name}\". If you do not see these files available for you to download, please let me know.")
    doc.add_paragraph("")
    doc.add_paragraph(f"Exam Schedule - {exam_date} at {exam_time}")
    doc.add_paragraph("")
    doc.add_paragraph(f"{mcq_start_time} - {mcq_end_time}: {exam_name} MCQs (100 questions; {mcq_duration})")
    doc.add_paragraph(f"{essay_start_time} - {essay_end_time}: {exam_name} Essay (1 question; {essay_duration})")
    doc.add_paragraph("Your exam must be taken in the assigned examination room at the assigned time, unless you have been given other instructions by OMESA staff.")
    doc.add_paragraph("You can take breaks during the exam. However, laptops must remain in the testing room until the end of the testing day, including during breaks in the middle of AND BETWEEN exam components.")
    doc.add_paragraph("You will have the opportunity to review the MCQ results immediately after completing the MCQ exam (although this is not required).")
    doc.add_paragraph("")
    doc.add_paragraph(f"LOCATION: {exam_location}")
    doc.add_paragraph("")
    doc.add_paragraph("Test Day Rules and Procedures")
    doc.add_paragraph("")
    doc.add_paragraph("Please arrive five minutes prior to the start of each exam or exam section.")
    doc.add_paragraph("Do not bring any personal/unauthorized items into the secure testing area. Such items include but are not limited to outerwear, hats, food, drinks, purses, briefcases, notebooks, notes, pagers, watches, cell phones, recording devices, and photographic equipment.")
    doc.add_paragraph("Food and drinks are only allowed if their presence is an approved testing accommodation by the Department of Accessibility Services.")
    doc.add_paragraph("You are not permitted to access any unauthorized items during exam administration.")
    doc.add_paragraph("A scratch pad or paper will be provided to you. You are not allowed to bring your own paper into the secure testing area.")
    doc.add_paragraph("Do not make notes on your scratch paper prior to starting your exam and/or entering your start-up code. Once your exam begins, you are permitted to make calculations or notes ONLY on the erasable note board or scratch paper provided.")
    doc.add_paragraph("You must turn in all used and unused scratch paper to the proctor at the end of your exam.")
    doc.add_paragraph("You must adhere to the instructions provided by proctors administering the examination.")
    doc.add_paragraph("Carefully review and agree to abide by any instructions provided or that appear at the start of the examination session.")
    doc.add_paragraph("While you may step away from your computer during an exam, you are NOT permitted to reference any outside materials until after that exam has been completed.")
    doc.add_paragraph("Test proctors are not authorized to answer questions from examinees regarding examination content or scoring during the exam.")
    doc.add_paragraph("Earplugs and ear protectors are permitted, though AirPods and other sound-playing devices are not.")
    doc.add_paragraph("Exams are closely monitored. Staff may be stationed in or around the testing area during the examination.")
    doc.add_paragraph("If cheating occurs, an affected individual’s exam may be stopped, the course director and dean will be notified, and all materials will be held for action.")
    doc.add_paragraph("")
    doc.add_paragraph("Reminders")
    doc.add_paragraph("")
    doc.add_paragraph("Make sure you close and save all your programs before opening ExamSoft.")
    doc.add_paragraph("Turn off any antivirus applications you have on your computer before opening ExamSoft. Please make sure you have the latest version of ExamSoft installed.")
    doc.add_paragraph("When you finish uploading/submitting your exam, ExamSoft will sometimes change the wallpaper on your desktop. You may need to change it back.")
    doc.add_paragraph("Please make sure that your completed exam uploads back to the server before powering down your computer for the day.")
    doc.add_paragraph("In order to permit same-day test reviews, essay examinations must be taken first, followed by multiple-choice examinations. Even if you do not intend to avail yourself of the test review, you are still required to take the essay examination first, unless specifically instructed not to do so.")
    doc.add_paragraph("Referencing notes during scheduled breaks between different components of an exam (MCQ, oral, essay, etc.) is permitted, but this is not true for breaks you take in the middle of an exam component. However, your computer must stay in the room at all times, even during breaks.")
    doc.add_paragraph("Scratch paper should not leave the exam room; please leave it at the front of the room before you depart for the day.")
    doc.add_paragraph("")
    doc.add_paragraph("Should you have any questions or concerns, please let me know!")

    # Save the document
    file_path = r"C:/Users/user_name/Documents/Exam_Notification.docx"
    doc.save(file_path)

    # Inform the user
    messagebox.showinfo("Success", "The exam notification has been generated and saved as 'Exam_Notification.docx'.")

# Run the function to generate the email
generate_exam_email()


 
