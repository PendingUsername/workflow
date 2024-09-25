import tkinter as tk
from tkinter import simpledialog, messagebox
from docx import Document

def generate_grade_notification():
    # Create a Tkinter root window and hide it
    root = tk.Tk()
    root.withdraw()

    # Prompt the user for input
    course_name = simpledialog.askstring("Input", "Enter the course name (e.g., EUSOM-MD 602 - FDNS: Female Reproduction):", parent=root)
    irat_percent = simpledialog.askstring("Input", "Enter the IRAT percentage:", parent=root)
    grat_percent = simpledialog.askstring("Input", "Enter the GRAT percentage:", parent=root)
    mcq_percent = simpledialog.askstring("Input", "Enter the MCQ percentage:", parent=root)
    essay_percent = simpledialog.askstring("Input", "Enter the Essay percentage:", parent=root)
    final_grade_percent = simpledialog.askstring("Input", "Enter the Final Grade percentage:", parent=root)
    grade_class_avg = simpledialog.askstring("Input", "Enter the Grade Class Average (e.g., 86.12 %):", parent=root)
    grade_std_dev = simpledialog.askstring("Input", "Enter the Grade Standard Deviation (e.g., 4.98):", parent=root)
    exam_avg = simpledialog.askstring("Input", "Enter the Final Exam Average (e.g., 86 %):", parent=root)
    exam_std_dev = simpledialog.askstring("Input", "Enter the Exam Standard Deviation (e.g., 5.09):", parent=root)

    # Create a Word document
    doc = Document()

    # Add content to the document based on user input
    doc.add_paragraph("Dear Students,")
    doc.add_paragraph("")
    doc.add_paragraph(f"We hope this message finds you well. We wanted to inform you that the grades for {course_name} have been posted and are now available for viewing.")
    doc.add_paragraph("")
    doc.add_paragraph("We encourage you to log in to the student portal or the designated grade management system to access your individual grades. Please take a moment to review your performance.")
    doc.add_paragraph("")
    doc.add_paragraph("Should you have any questions or concerns regarding your grades or any other course-related matters, please don't hesitate to reach out to Dr. Schulman or Mr. Cook, our assessment coordinator. They will be more than happy to assist you and provide any clarification you may need.")
    doc.add_paragraph("")
    doc.add_paragraph(f"Congratulations on completing the {course_name} course. We appreciate your dedication and hard work throughout the term. Keep up the excellent effort as you continue your academic journey.")
    doc.add_paragraph("")
    doc.add_paragraph(f"Your {course_name} exam scores as well as your final Satisfactory/Unsatisfactory (S/U) block grade have been posted to OASIS.")
    doc.add_paragraph("")
    doc.add_paragraph("You will see the following when you log in OASIS:")
    doc.add_paragraph(f"IRAT: {irat_percent}")
    doc.add_paragraph(f"GRAT: {grat_percent}")
    doc.add_paragraph(f"MCQ: {mcq_percent}")
    doc.add_paragraph(f"Essay: {essay_percent}")
    doc.add_paragraph(f"Final Grade: {final_grade_percent}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Exam Stats & Notes:")
    doc.add_paragraph("")
    doc.add_paragraph(f"{course_name}")
    doc.add_paragraph(f"Grade Class Average: {grade_class_avg}")
    doc.add_paragraph(f"Standard Deviation: {grade_std_dev}")
    doc.add_paragraph(f"Final Exam Average: {exam_avg}")
    doc.add_paragraph(f"Standard Deviation: {exam_std_dev}")
    doc.add_paragraph("")
    doc.add_paragraph("If you have any questions, please feel free to reach out!")

    # Save the document to the specified location
    file_path = r"C:\Users\user_name\Documents\Grade_Notification.docx"
    doc.save(file_path)

    # Inform the user
    messagebox.showinfo("Success", f"The grade notification has been generated and saved as '{file_path}'.")

# Run the function to generate the email
generate_grade_notification()
